import os, json
from langchain.schema import Document
from docx import Document as DocxDocument
import xml.etree.ElementTree as ET
from langchain_community.document_loaders import UnstructuredFileLoader, TextLoader

def load_docx(path: str):
    doc = DocxDocument(path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return Document(page_content='\n'.join(texts), metadata={'source': path})

def load_xml(path: str):
    tree = ET.parse(path)
    root = tree.getroot()
    texts = []
    for elem in root.iter():
        if elem.text and elem.text.strip():
            texts.append(elem.text.strip())
    return Document(page_content='\n'.join(texts), metadata={'source': path})

def load_documents(path: str):
    """Load documents. If path is directory, load all supported files inside."""
    docs = []
    if os.path.isdir(path):
        for fname in sorted(os.listdir(path)):
            fpath = os.path.join(path, fname)
            if os.path.isfile(fpath):
                lname = fname.lower()
                if lname.endswith('.docx') or lname.endswith('.doc'):
                    docs.append(load_docx(fpath))
                elif lname.endswith('.xml'):
                    docs.append(load_xml(fpath))
                elif lname.endswith('.json'):
                    try:
                        with open(fpath, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        if isinstance(data, list):
                            for item in data:
                                text = item.get('text') or item.get('content') or json.dumps(item, ensure_ascii=False)
                                docs.append(Document(page_content=text, metadata={'source': fpath}))
                        else:
                            docs.append(Document(page_content=json.dumps(data, ensure_ascii=False), metadata={'source': fpath}))
                    except Exception:
                        try:
                            loader = TextLoader(fpath, encoding='utf-8')
                            docs.extend(loader.load())
                        except Exception:
                            pass
                else:
                    try:
                        loader = UnstructuredFileLoader(fpath)
                        docs.extend(loader.load())
                    except Exception:
                        pass
    elif os.path.isfile(path):
        lname = path.lower()
        if lname.endswith('.docx') or lname.endswith('.doc'):
            return [load_docx(path)]
        elif lname.endswith('.xml'):
            return [load_xml(path)]
    else:
        raise FileNotFoundError(path)
    return docs
